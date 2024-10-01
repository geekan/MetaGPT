from pydantic import BaseModel
from typing import List, Union, Any
from pathlib import Path
import re
from datetime import datetime
from docx import Document
 
def colored_decorator(color):   
    def print_colored(func):  
        async def wrapper(self, *args, **kwargs):
                # ANSI escape codes for colored terminal output (optional, can be removed if not needed)
                print(color)
                contexts = await func(self, *args, **kwargs)  
                print("\033[0m")
                return contexts             
        return wrapper
    return print_colored


def print_time(func):  
    async def wrapper(self, *args, **kwargs):
            # ANSI escape codes for colored terminal output (optional, can be removed if not needed)
            start_time = datetime.now()
            objs = await func(self, *args, **kwargs)  
            end_time = datetime.now()
            print(f"执行时间：{(end_time - start_time).total_seconds()}秒")
            return objs             
    return wrapper


class Node:
        def __init__(self, name='',content = ''):
            self.name = name
            self.content = content
            self.subheading = {}
            
class TitleHierarchy:
    def __init__(self):
        self.root = Node()
    
    def add_chapter(self, id: str, name: str):
        # Create a new node for the chapter
        chapter_node = Node(name)
        # Add the new chapter node to the root's subheading
        self.root.subheading[id] = chapter_node
    
    
    def add_title(self, title: str):
        parts = title.split(' ')
        path = parts[0].split('.')
        name = ' '.join(parts[1:])
        current = self.root

        for part in path:
            if part not in current.subheading:
                current.subheading[part] = Node()
            current = current.subheading[part]

        current.name = name
    
    def add_subheadings(self, titles):
        """
        Add multiple subheadings to the hierarchy.

        param titles: A list of title strings to be added to the hierarchy.
        """
        for title in titles:
            self.add_title(title)
    
    
    def get_subheadings_by_prefix(self, prefix):
        """
        Get subheadings based on the given prefix.

        param root: The root Node of the hierarchical structure representing the relationship between titles and subtitles.
        param prefix: The prefix string to filter the subheadings.
        return: A list of subheading names that match the given prefix.
        """
        # Split the prefix into its components
        prefix_parts = prefix.split('.')

        # Navigate through the structure using the prefix
        current = self.root
        for part in prefix_parts:
            if part in current.subheading:
                current = current.subheading[part]
            else:
                # If any part of the prefix is not found, return an empty list
                return []

        # Collect all subheading names at this level
        subheadings = []
        for key, value in current.subheading.items():
            subheadings.append(f"{prefix}.{key} {value.name}")

        return subheadings

    def set_content_by_id(self, path: str, content: str):
        """
        Set the content of a node identified by its path.

        param path: A string representing the path to the node, formatted as "x.y.z...".
        param content: The content to be set for the node.
        """
        # Split the path into its components
        path_parts = path.split('.')

        # Navigate through the structure using the path
        current = self.root
        for part in path_parts:
            if part in current.subheading:
                current = current.subheading[part]
            else:
                # If any part of the path does not exist, return without setting content
                return

        # Set the content of the node
        current.content = content 
        
    def set_content_by_headings(self, titles, contents):
        """
        Set the contents of nodes identified by their paths.

        param titles: A list of title strings representing the paths.
        param contents: A list of contents corresponding to the titles.
        """
        for title, content in zip(titles, contents):
            path, _  =  title.split(' ')
            self.set_content_by_id(path, content)

    def traverse_and_output(self, node=None, prefix='', level=0):
        if node is None:
            node = self.root

        output = []
        for key, child in node.subheading.items():
            title = f"{prefix}{key} {child.name}" if prefix else f"{key} {child.name}"
            output.append((title, child.content, level + 1))
            output.extend(self.traverse_and_output(child, f"{prefix}{key}.", level + 1))

        return output
    
    def get_chapter_obj_by_id(self, id: str) -> str:
        """
        Get the chapter name by its hierarchical ID.

        param id: The hierarchical ID of the chapter to look up, formatted as "x.y.z...".
        return: The name of the chapter if found, otherwise an empty string.
        """
        # Split the hierarchical ID into its components
        id_parts = id.split('.')
        # Start from the root and navigate through the structure using the ID
        current = self.root
        for part in id_parts:
            if part in current.subheading:
                current = current.subheading[part]
            else:
                # If any part of the ID is not found, return an empty string
                return ''
        # Return the name of the chapter node
        return current
    


class WriteOutFile:
    
    @staticmethod
    def write_markdown_file(topic: str, tasks: Any, output_path: Union[str, Path]):
        pass
    
    @staticmethod
    def write_word_file(topic: str, tasks: Any, output_path: Union[str, Path]):
        """
        Writes tasks to a Word document.

        topic (str): The main topic of the document.
        tasks (List[Tuple[str, str, int]]): A list of tuples containing the title, content, and heading level of each task.
        output_path (Union[str, Path]): The file path where the document will be saved.
        """
        def post_processes(title: str, context: str) -> str:
            """Post-processes the context by removing the first line that match the title."""
            normalize_text = lambda x:  re.sub(r'[^\u4e00-\u9fa5]+', '', x)
            split_context = context.split('\n') if context else []
            prefix = split_context[0] if split_context else ''
            if  normalize_text(title) == normalize_text(prefix): 
                context = '\n'.join(split_context[1:])
            return context  
        
        # Ensure the write_out_file is a Path object
        output_path = Path(output_path)
        document = Document()
        document.add_heading(topic, level=0)
        
        # Process each task
        for title, content, level  in tasks:
            # Add a heading for the task
            document.add_heading(title, level=level)
            # Post-process and add it to the document
            content = post_processes(title, content)
            document.add_paragraph(content)
            
        document.add_page_break()
        # Save the document
        try:
            document.save(output_path)
        except Exception as e:
            print(f"An error occurred while saving the document: {e}")   
    
  
       
         